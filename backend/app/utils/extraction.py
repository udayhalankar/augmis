from pathlib import Path
import logging
import shutil
import subprocess
import pandas as pd
from pypdf import PdfReader
from docx import Document

from app.core.config import settings

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None

try:
    import pypdfium2 as pdfium
except Exception:  # pragma: no cover - optional dependency
    pdfium = None


SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".csv", ".txt", ".docx"}
PDF_TEXT_MIN_CHARS_FOR_NON_OCR = 80
COMMON_TESSERACT_PATHS = [
    "/usr/bin/tesseract",
    "/usr/local/bin/tesseract",
    "C:/Program Files/Tesseract-OCR/tesseract.exe",
    "C:/Program Files (x86)/Tesseract-OCR/tesseract.exe",
    "C:/Users/udayh/AppData/Local/Programs/Tesseract-OCR/tesseract.exe",
]
logger = logging.getLogger(__name__)


def _is_valid_tesseract_candidate(candidate: str) -> bool:
    try:
        completed = subprocess.run(
            [candidate, "--version"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
            timeout=5,
        )
        return completed.returncode == 0
    except Exception:
        return False


def _resolve_tesseract_cmd() -> str | None:
    configured_path = str(settings.OCR_TESSERACT_CMD or "").strip()
    candidates = [configured_path] if configured_path else []

    path_binary = shutil.which("tesseract")
    if path_binary:
        candidates.append(path_binary)

    candidates.extend(COMMON_TESSERACT_PATHS)

    seen: set[str] = set()
    for candidate in candidates:
        if not candidate:
            continue
        normalized = str(candidate).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        if Path(normalized).exists() and _is_valid_tesseract_candidate(normalized):
            return normalized
        if Path(normalized).name == normalized and _is_valid_tesseract_candidate(normalized):
            return normalized

    return None


def get_ocr_diagnostics() -> dict:
    tesseract_cmd = _resolve_tesseract_cmd()
    available, error = _ocr_is_available()
    return {
        "available": available,
        "error": error,
        "tesseract_cmd": tesseract_cmd,
        "configured_tesseract_cmd": str(settings.OCR_TESSERACT_CMD or "").strip() or None,
        "pytesseract_installed": pytesseract is not None,
        "pypdfium2_installed": pdfium is not None,
    }


def _ocr_is_available() -> tuple[bool, str | None]:
    if pytesseract is None:
        return False, "pytesseract is not installed"

    if pdfium is None:
        return False, "pypdfium2 is not installed"

    tesseract_cmd = _resolve_tesseract_cmd()
    if not tesseract_cmd:
        configured_path = str(settings.OCR_TESSERACT_CMD or "").strip()
        if configured_path:
            return False, "configured tesseract binary is unavailable or failed validation"
        return False, "tesseract binary is not installed or not on PATH"

    try:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    except Exception:
        return False, "tesseract binary is configured but could not be initialized"

    return True, None


def extract_pdf(path: Path) -> str:
    return extract_pdf_with_details(path)["text"]


def _extract_pdf_text_layer(path: Path) -> dict:
    text_parts = []
    reader = PdfReader(str(path))

    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
            text_parts.append(f"[Page {i + 1}] {page_text}")
        except Exception as exc:
            text_parts.append(f"[Page {i + 1}] Error extracting page: {exc}")

    full_text = "\n".join(text_parts)
    return {
        "text": full_text,
        "page_count": len(reader.pages),
        "extracted_characters": len(full_text.strip()),
    }


def _extract_pdf_with_ocr(path: Path) -> dict:
    available, error = _ocr_is_available()
    if not available:
        logger.warning(
            "OCR unavailable for PDF extraction",
            extra={
                "category": "ocr_unavailable",
                "metadata": {
                    "file_name": path.name,
                    "ocr_error": error,
                },
            },
        )
        return {
            "text": "",
            "ocr_used": False,
            "ocr_available": False,
            "ocr_error": error,
            "extracted_characters": 0,
        }

    pdf = pdfium.PdfDocument(str(path))
    text_parts = []

    for index in range(len(pdf)):
        page = pdf[index]
        bitmap = page.render(scale=2.0)
        pil_image = bitmap.to_pil()
        try:
            page_text = pytesseract.image_to_string(pil_image) or ""
        except Exception as exc:
            logger.exception(
                "OCR page extraction failed",
                extra={
                    "category": "ocr_extraction_failure",
                    "is_critical": True,
                    "metadata": {
                        "file_name": path.name,
                        "page_number": index + 1,
                    },
                },
            )
            page_text = f"[OCR error: {exc}]"
        text_parts.append(f"[Page {index + 1}] {page_text}")

    full_text = "\n".join(text_parts)
    return {
        "text": full_text,
        "ocr_used": True,
        "ocr_available": True,
        "ocr_error": None,
        "extracted_characters": len(full_text.strip()),
    }


def extract_pdf_with_details(path: Path) -> dict:
    text_layer = _extract_pdf_text_layer(path)
    text_layer_chars = text_layer["extracted_characters"]

    if text_layer_chars >= PDF_TEXT_MIN_CHARS_FOR_NON_OCR:
        return {
            "text": text_layer["text"],
            "parser": "pypdf",
            "ocr_used": False,
            "ocr_available": _ocr_is_available()[0],
            "ocr_error": None,
            "page_count": text_layer["page_count"],
            "extracted_characters": text_layer_chars,
            "text_status": "text_layer",
        }

    ocr_result = _extract_pdf_with_ocr(path)
    if ocr_result["extracted_characters"] > text_layer_chars:
        return {
            "text": ocr_result["text"],
            "parser": "tesseract_ocr",
            "ocr_used": ocr_result["ocr_used"],
            "ocr_available": ocr_result["ocr_available"],
            "ocr_error": ocr_result["ocr_error"],
            "page_count": text_layer["page_count"],
            "extracted_characters": ocr_result["extracted_characters"],
            "text_status": "ocr_text",
        }

    return {
        "text": text_layer["text"],
        "parser": "pypdf",
        "ocr_used": False,
        "ocr_available": ocr_result["ocr_available"],
        "ocr_error": ocr_result["ocr_error"],
        "page_count": text_layer["page_count"],
        "extracted_characters": text_layer_chars,
        "text_status": "low_text" if text_layer_chars > 0 else "empty_text",
    }


def extract_excel(path: Path) -> str:
    text_parts = []
    xl = pd.ExcelFile(path)

    for sheet in xl.sheet_names:
        df = xl.parse(sheet).fillna("")
        text_parts.append(f"\nWorkbook: {path.name} | Sheet: {sheet}\n")
        text_parts.append(df.to_csv(index=False))

        for idx, row in df.iterrows():
            row_text = "; ".join(f"{col}: {row[col]}" for col in df.columns)
            text_parts.append(f"Row {idx + 2}: {row_text}")

    return "\n".join(text_parts)


def extract_csv(path: Path) -> str:
    df = pd.read_csv(path).fillna("")
    text_parts = [df.to_csv(index=False)]

    for idx, row in df.iterrows():
        row_text = "; ".join(f"{col}: {row[col]}" for col in df.columns)
        text_parts.append(f"Row {idx + 2}: {row_text}")

    return "\n".join(text_parts)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts)


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_file_text_with_details(path: Path) -> dict:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_with_details(path)

    if suffix in {".xlsx", ".xls"}:
        text = extract_excel(path)
        return {
            "text": text,
            "parser": "pandas_excel",
            "ocr_used": False,
            "ocr_available": _ocr_is_available()[0],
            "ocr_error": None,
            "page_count": None,
            "extracted_characters": len(text.strip()),
            "text_status": "parsed" if text.strip() else "empty_text",
        }

    if suffix == ".csv":
        text = extract_csv(path)
        return {
            "text": text,
            "parser": "pandas_csv",
            "ocr_used": False,
            "ocr_available": _ocr_is_available()[0],
            "ocr_error": None,
            "page_count": None,
            "extracted_characters": len(text.strip()),
            "text_status": "parsed" if text.strip() else "empty_text",
        }

    if suffix == ".docx":
        text = extract_docx(path)
        return {
            "text": text,
            "parser": "python_docx",
            "ocr_used": False,
            "ocr_available": _ocr_is_available()[0],
            "ocr_error": None,
            "page_count": None,
            "extracted_characters": len(text.strip()),
            "text_status": "parsed" if text.strip() else "empty_text",
        }

    if suffix == ".txt":
        text = extract_txt(path)
        return {
            "text": text,
            "parser": "plain_text",
            "ocr_used": False,
            "ocr_available": _ocr_is_available()[0],
            "ocr_error": None,
            "page_count": None,
            "extracted_characters": len(text.strip()),
            "text_status": "parsed" if text.strip() else "empty_text",
        }

    return {
        "text": "",
        "parser": "unsupported",
        "ocr_used": False,
        "ocr_available": _ocr_is_available()[0],
        "ocr_error": None,
        "page_count": None,
        "extracted_characters": 0,
        "text_status": "unsupported",
    }


def extract_file_text(path: Path) -> str:
    return extract_file_text_with_details(path)["text"]


def scan_documents(datasource: str) -> list[Path]:
    root = Path(datasource)

    if not root.exists():
        return []

    return [
        p for p in root.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
